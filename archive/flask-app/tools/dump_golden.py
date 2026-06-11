# dump_golden.py — one-time golden/parity data generator for the JS port.
#
# Runs the CURRENT Python classifier over every corpus text and emits:
#   test/corpus/*.json        committed corpus cases (texts + expectations)
#   test/corpus-local/*.json  professor-material corpus (gitignored)
#   test/golden/python-golden.json  full classify_text output per case
#
# Golden output for committed corpus cases goes to test/golden/ (committed);
# golden output for professor-material cases goes to test/corpus-local/
# (gitignored) so no exam content ever reaches the public repo.
#
# Run from the repo root:
#   python archive/flask-app/tools/dump_golden.py
import json
import os
import re
import sys

FLASK_APP_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, FLASK_APP_DIR)

from docx import Document  # noqa: E402
from app.classifier import classify_text  # noqa: E402
from app.parser import parse_plaintext  # noqa: E402

ROOT = os.path.dirname(os.path.dirname(FLASK_APP_DIR))
TEST_DOCS = os.path.join(ROOT, "archive", "test-docs")


def slug(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


# ---------------------------------------------------------------------------
# Inline pytest corpus
# ---------------------------------------------------------------------------

SPECTRUM_CASES = [
    {
        "name": "full issue-led IRAC",
        "text": (
            "The issue is whether Dana committed battery. "
            "Battery requires an intentional harmful or offensive contact with another person. "
            "Dana shoved Lee in the shoulder during the argument, and Lee stumbled backward. "
            "Therefore, Dana is likely liable for battery."
        ),
        "framework": "AUTO",
        "expect": {"badge": "COMPLETE_IRAC", "labels": ["ISSUE", "RULE", "APPLICATION", "CONCLUSION"]},
    },
    {
        "name": "issue-less RAC should not be called IRAC",
        "text": (
            "Battery requires an intentional harmful or offensive contact with another person. "
            "Dana shoved Lee in the shoulder during the argument, and Lee stumbled backward. "
            "Therefore, Dana is likely liable for battery."
        ),
        "framework": "AUTO",
        "expect": {"badge": "COMPLETE_RAC", "labels": ["RULE", "APPLICATION", "CONCLUSION"]},
    },
    {
        "name": "automatic CRAC",
        "text": (
            "Dana is likely liable for battery. "
            "Battery requires an intentional harmful or offensive contact with another person. "
            "Dana shoved Lee in the shoulder during the argument, and Lee stumbled backward. "
            "Therefore, Dana is likely liable."
        ),
        "framework": "AUTO",
        "expect": {"badge": "COMPLETE_CRAC", "effective_framework": "CRAC"},
    },
    {
        "name": "automatic CREAC",
        "text": (
            "The covenant is likely enforceable against Garrison. "
            "An equitable servitude binds successors when the covenant touches and concerns land, the parties intend it to run, and the later owner has notice. "
            "In Neponsit, the court enforced a covenant because the promise affected property ownership and bound successors with notice. "
            "Here, Garrison bought Parcel 14 after the residential-use restriction was recorded in the chain of title. "
            "Accordingly, the covenant is likely enforceable against Garrison."
        ),
        "framework": "AUTO",
        "expect": {"badge": "COMPLETE_CREAC", "effective_framework": "CREAC"},
    },
    {
        "name": "missing application in rule dump",
        "text": (
            "The issue is whether the agreement is enforceable. "
            "A valid contract requires offer, acceptance, and consideration. "
            "Courts apply the objective theory of assent."
        ),
        "framework": "AUTO",
        "expect": {"badge": "WARNING_MISSING_APPLICATION"},
    },
    {
        "name": "muddled blend",
        "text": (
            "The issue is whether Rivera breached the duty of care. "
            "A driver must use reasonable care, and here Rivera looked down at a text message for six seconds while entering the intersection. "
            "Therefore, Rivera likely breached the duty of care."
        ),
        "framework": "AUTO",
        "expect": {"badge": "WARNING_BLEND"},
    },
    {
        "name": "NeueStruktur general jurisdiction application",
        "text": (
            "Whether the Western District of Virginia may exercise general personal jurisdiction over NeueStruktur GmbH, a German corporation, based on its contacts with Virginia. "
            "Under Daimler AG v. Bauman, 571 U.S. 117 (2014), and Goodyear Dunlop Tires Operations, S.A. v. Brown, 564 U.S. 915 (2011), a court may exercise general personal jurisdiction over a corporation only where the corporation's affiliations with the forum are so continuous and systematic as to render it 'essentially at home' there. "
            "For corporations, the paradigm cases of general jurisdiction are the state of incorporation and the state of principal place of business. "
            "Only in an 'exceptional case' would operations in another forum be so substantial as to render the corporation at home there. "
            "Even substantial and continuous commercial activity in a forum does not make a corporation 'at home' where it is neither incorporated nor headquartered. "
            "NeueStruktur GmbH is incorporated under German law and headquartered in Frankfurt, Germany. "
            "It has no offices, employees, agents, or bank accounts in Virginia or anywhere in the United States. "
            "Its U.S. subsidiary (NSA) is a separate legal entity with a New York principal place of business; NSA's contacts cannot automatically be imputed to the German parent absent a formal agency finding. "
            "Even if NSA's contacts were fully imputed, NSA is at home in Delaware and New York — not Virginia. "
            "NeueStruktur's five site visits to Virginia over two years are episodic, not continuous and systematic. "
            "Virginia is not among the paradigm forums for NeueStruktur, and no exceptional circumstances exist to support general jurisdiction there. "
            "The Western District of Virginia cannot exercise general personal jurisdiction over NeueStruktur GmbH."
        ),
        "framework": "AUTO",
        "expect": {
            "badge": "COMPLETE_IRAC",
            "must_label": {
                "Virginia is not among the paradigm forums": "APPLICATION",
                "Even substantial and continuous commercial activity": "RULE",
            },
        },
    },
]

_FRAUD_BASE = (
    "The issue is whether Mercer committed fraud. "
    "To establish fraud, a plaintiff must prove that the defendant made a false representation, "
    "knew it was false, intended reliance, and caused damages. "
)

PERTURBATION_CASES = [
    {
        "name": "baseline with signposts",
        "text": _FRAUD_BASE
        + "Here, Mercer told Okafor that the fund returned twelve percent annually, but Mercer knew this was false. "
        + "Therefore, Okafor is likely to prevail.",
    },
    {
        "name": "removed Here",
        "text": _FRAUD_BASE
        + "Mercer told Okafor that the fund returned twelve percent annually, but Mercer knew this was false. "
        + "Therefore, Okafor is likely to prevail.",
    },
    {
        "name": "removed Therefore",
        "text": _FRAUD_BASE
        + "Here, Mercer told Okafor that the fund returned twelve percent annually, but Mercer knew this was false. "
        + "Okafor is likely to prevail.",
    },
    {
        "name": "removed both signposts",
        "text": _FRAUD_BASE
        + "Mercer told Okafor that the fund returned twelve percent annually, but Mercer knew this was false. "
        + "Okafor is likely to prevail.",
    },
    {
        "name": "citation-heavy application without Here",
        "text": (
            "The issue is whether transfer was proper. "
            "Under 28 U.S.C. § 1404(a), a court may transfer for convenience when venue is proper. "
            "The case was transferred under § 1406(a), but the original venue was proper under § 1391. "
            "The transfer was improper."
        ),
    },
]
for case in PERTURBATION_CASES:
    case["framework"] = "IRAC"
    case["expect"] = {
        "badge": "COMPLETE_IRAC",
        "labels": ["ISSUE", "RULE", "APPLICATION", "CONCLUSION"],
    }

COUNTERARGUMENT_CASES = [
    {
        "name": "counterargument gets coaching metadata",
        "text": (
            "The issue is whether Dana committed battery. "
            "Battery requires intentional harmful or offensive contact. "
            "The defendant will argue that the contact was accidental. "
            "However, Dana raised her hand and shoved Lee after the warning. "
            "Therefore, Dana is likely liable."
        ),
        "framework": "AUTO",
        "expect": {
            "must_label": {"defendant will argue": "APPLICATION"},
            "must_counterargument": ["defendant will argue"],
            "priority_kinds": ["counterargument_needs_response"],
        },
    },
    {
        "name": "counterargument is not bare conclusion in CRAC opener",
        "text": (
            "Defendant will argue that the search was reasonable. "
            "The Fourth Amendment requires a warrant unless an exception applies. "
            "Here, the officer searched the bag without consent or exigency. "
            "Therefore, the search was likely unreasonable."
        ),
        "framework": "CRAC",
        "expect": {
            "must_label": {"Defendant will argue": "APPLICATION"},
            "must_counterargument": ["Defendant will argue"],
        },
    },
    {
        "name": "citation-heavy application stays application",
        "text": (
            "The issue is whether transfer was proper. "
            "Under 28 U.S.C. § 1404(a), a court may transfer for convenience when venue is proper. "
            "The case was transferred under § 1406(a), but the original venue was proper under § 1391. "
            "Therefore, the transfer was improper."
        ),
        "framework": "IRAC",
        "expect": {"must_label": {"transferred under": "APPLICATION"}},
    },
]

TRAINING_CASES = [
    {
        "name": "sentence diagnostics",
        "text": "The court should consider fairness and efficiency.",
        "framework": "AUTO",
        "expect": {"first_sentence_confidence_label_in": ["low", "medium"]},
    },
    {
        "name": "paragraph training fields",
        "text": "Here, Dana shoved Lee during the argument. Therefore, Dana is likely liable.",
        "framework": "AUTO",
        "expect": {"priority_kinds": ["missing_rule", "unsupported_application"]},
    },
    {
        "name": "creac explanation and competing labels",
        "text": (
            "Dana is likely liable for battery. "
            "Battery requires intentional harmful or offensive contact. "
            "In Garratt v. Dailey, the court found that intent existed because contact was substantially certain. "
            "Here, Dana shoved Lee in the shoulder. "
            "Therefore, Dana is likely liable."
        ),
        "framework": "AUTO",
        "expect": {"effective_framework": "CREAC", "has_label": ["EXPLANATION"]},
    },
    {
        "name": "cross paragraph training priority",
        "paragraphs": [
            "Battery requires intentional harmful or offensive contact.",
            "Here, Dana shoved Lee during the argument. Therefore, Dana is likely liable.",
        ],
        "framework": "AUTO",
        "expect": {"priority_kinds_para0": ["cross_paragraph_split"]},
    },
]


# ---------------------------------------------------------------------------
# Committed .docx corpus extraction
# ---------------------------------------------------------------------------

def load_structure_doc_cases(docx_path):
    """Test Document v1.0 — returns {case_id: (framework, joined_text)}."""
    doc = Document(docx_path)
    cases = {}
    current_id = None
    current_fw = "IRAC"
    current_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m = re.match(r"^TEST CASE ([ABC]-\d+)", text)
        if m:
            if current_id:
                cases[current_id] = (current_fw, " ".join(current_paras))
            current_id = m.group(1)
            current_fw = "CREAC" if "CREAC" in text.upper() else "IRAC"
            current_paras = []
            continue
        if "PART D" in text or "Answer Key" in text:
            break
        if text.startswith("Note:"):
            continue
        if re.match(r"^PART [A-Z]", text):
            continue
        if re.match(r"^(This document|Color legend|Prepared for)", text):
            continue
        if re.match(r"^IRAC/CREAC Structural", text):
            continue
        if re.match(r"^Test Document", text):
            continue
        if current_id:
            current_paras.append(text)
    if current_id and current_paras:
        cases[current_id] = (current_fw, " ".join(current_paras))
    return cases


STRUCTURE_DOC_EXPECTED = {
    "A-1": ("IRAC", "COMPLETE_IRAC"),
    "A-2": ("IRAC", "COMPLETE_IRAC"),
    "A-3": ("IRAC", "COMPLETE_IRAC"),
    "B-1": ("CREAC", "COMPLETE_CREAC"),
    "B-2": ("CREAC", "COMPLETE_CREAC"),
    "C-1": ("IRAC", "WARNING_MISSING_RULE"),
    "C-2": ("IRAC", "WARNING_MISSING_APPLICATION"),
    "C-3": ("IRAC", "WARNING_MISSING_CONCLUSION"),
    "C-4": ("IRAC", "WARNING_BLEND"),
    "C-5": ("IRAC", "WARNING_APPLICATION_WITHOUT_RULE"),
    "C-6": ("IRAC", "INFO_MOSTLY_UNCLASSIFIED"),
    "C-7": ("CREAC", "WARNING_EXPLANATION_AFTER_APPLICATION"),
}

EC_PARA_INDICES = {
    "EC-01": ([18, 20, 22, 24, 26], "IRAC"),
    "EC-02": ([60, 62, 64, 66], "IRAC"),
    "EC-03": ([96], "IRAC"),
    "EC-04": ([102], "IRAC"),
    "EC-05": ([72, 74], "IRAC"),
    "EC-06": ([124, 126, 128, 130], "IRAC"),
    "EC-07": ([32, 34, 36, 38], "IRAC"),
    "EC-08": ([108], "IRAC"),
    "EC-09": ([136, 138, 140, 142, 144], "CREAC"),
    "EC-10": ([114], "IRAC"),
    "EC-11": ([80, 82, 84, 86], "IRAC"),
    "EC-12": ([44, 46, 48, 50], "IRAC"),
}

EC_EXPECTED = {
    "EC-01": "COMPLETE_IRAC",
    "EC-02": "COMPLETE_IRAC",
    "EC-03": "INFO_INTRODUCTORY",
    "EC-04": "INFO_INTRODUCTORY",
    "EC-05": "WARNING_MISSING_APPLICATION",
    "EC-06": "COMPLETE_IRAC",
    "EC-07": "COMPLETE_IRAC",
    "EC-08": "COMPLETE_IRAC",
    "EC-09": "COMPLETE_CREAC",
    "EC-10": "COMPLETE_IRAC",
    "EC-11": "COMPLETE_IRAC",
    "EC-12": "COMPLETE_IRAC",
}


# ---------------------------------------------------------------------------
# Gitignored professor-material corpus (corpus-local)
# ---------------------------------------------------------------------------

MODEL_ANSWER_SECTIONS = [
    ("I-Q1", [9], "IRAC", "COMPLETE_RAC"),
    ("I-Q2", [15], "IRAC", "WARNING_APPLICATION_WITHOUT_RULE"),
    ("I-Q3", [19, 20, 21], "IRAC", "COMPLETE_RAC"),
    ("I-Q4", [25, 26, 27], "IRAC", "COMPLETE_RAC"),
    ("II-Q1", [32], "IRAC", "COMPLETE_RAC"),
    ("II-Q2", [36], "IRAC", "WARNING_MISSING_APPLICATION"),
    ("II-Q3", [42, 43], "IRAC", "WARNING_MISSING_RULE"),
    ("II-Q4", [47, 48, 49], "IRAC", "COMPLETE_RAC"),
    ("E1-S1", [56], "IRAC", "COMPLETE_IRAC"),
    ("E1-S2", [62], "IRAC", "COMPLETE_RAC"),
    ("E1-S3", [66, 67, 68, 69], "IRAC", "COMPLETE_IRAC"),
    ("E1-S4", [74], "IRAC", "COMPLETE_IRAC"),
    ("E1-S5", [78, 79, 80, 81], "IRAC", "COMPLETE_IRAC"),
    ("E2-S1", [87, 88, 89, 90], "IRAC", "COMPLETE_IRAC"),
    ("E2-S2", [94, 95, 96, 97], "IRAC", "COMPLETE_IRAC"),
    ("E2-S3", [103, 104, 105, 106], "IRAC", None),
    ("E2-S4", [111], "IRAC", "COMPLETE_IRAC"),
    ("E2-S5", [115], "IRAC", None),
    ("E2-S6", [119, 120, 121, 122], "IRAC", "COMPLETE_IRAC"),
]

# Cases whose badge expectation is unmet for a known, planned reason — flags
# flow into the corpus JSON as expect.pending and run as TODO in the JS suite.
# (Empty since the JS engine's cite-guarded application phrases landed; kept
# for future regeneration needs.)
MODEL_ANSWER_PENDING = {}

CIVPRO_SECTIONS = [
    ("P2-Q1", [155, 156, 157], "CRAC"),
    ("P2-Q2", [159, 160, 161], "CRAC"),
    ("P2-Q3", [163, 164, 165], "CRAC"),
    ("P2-Q4", [167, 168, 169], "CRAC"),
    ("E1-S1", [174, 175, 176, 177], "IRAC"),
    ("E1-S2", [179, 180, 181, 182, 183, 184], "IRAC"),
    ("E1-S3", [186, 187, 188, 189, 190, 191, 192], "IRAC"),
    ("E2-S1", [195, 196, 197, 198, 199, 200], "IRAC"),
    ("E2-S2", [202, 203, 204, 205], "IRAC"),
]

_CIVPRO_STRUCT_RE = re.compile(
    r"^(?P<lbl>issue|rule|analysis|application(?:\s*[—\-][^:]{0,60})?|conclusion)\s*:\s*",
    re.I,
)


def civpro_expected_label(para_text):
    m = _CIVPRO_STRUCT_RE.match(para_text)
    if not m:
        return None
    lbl = m.group("lbl").lower()
    if lbl.startswith("issue"):
        return "ISSUE"
    if lbl.startswith("rule"):
        return "RULE"
    if lbl.startswith("analysis") or lbl.startswith("application"):
        return "APPLICATION"
    if lbl.startswith("conclusion"):
        return "CONCLUSION"
    return None


# ---------------------------------------------------------------------------
# Build everything
# ---------------------------------------------------------------------------

def build_inline_file(source_name, cases):
    out = {"source": source_name, "cases": []}
    for case in cases:
        paragraphs = case.get("paragraphs") or [case["text"]]
        out["cases"].append({
            "id": f"{source_name}/{slug(case['name'])}",
            "framework": case["framework"],
            "paragraphs": paragraphs,
            "expect": case.get("expect", {}),
        })
    return out


def main():
    corpus_dir = os.path.join(ROOT, "test", "corpus")
    local_dir = os.path.join(ROOT, "test", "corpus-local")
    golden_dir = os.path.join(ROOT, "test", "golden")
    for d in (corpus_dir, local_dir, golden_dir):
        os.makedirs(d, exist_ok=True)

    corpus_files = {}

    corpus_files["spectrum.json"] = build_inline_file("spectrum", SPECTRUM_CASES)
    corpus_files["perturbations.json"] = build_inline_file("perturbations", PERTURBATION_CASES)
    corpus_files["counterarguments.json"] = build_inline_file("counterarguments", COUNTERARGUMENT_CASES)
    corpus_files["training.json"] = build_inline_file("training", TRAINING_CASES)

    # Structure doc v1.0 (committed)
    sd_cases = load_structure_doc_cases(
        os.path.join(TEST_DOCS, "Legal_Writing_Structure_Coach_Test_Document_v1.0.docx"))
    sd_out = {"source": "structure-doc-v1", "cases": []}
    for case_id in sorted(STRUCTURE_DOC_EXPECTED):
        fw_expected, badge = STRUCTURE_DOC_EXPECTED[case_id]
        fw, text = sd_cases[case_id]
        assert fw == fw_expected, f"{case_id}: framework mismatch {fw} != {fw_expected}"
        sd_out["cases"].append({
            "id": f"structure-doc-v1/{case_id}",
            "framework": fw,
            "paragraphs": [text],
            "expect": {"badge": badge},
        })
    corpus_files["structure-doc-v1.json"] = sd_out

    # Edge cases (committed)
    ec_doc = Document(os.path.join(TEST_DOCS, "Legal_Writing_Structure_Coach_EdgeCase_Test_v1.0.docx"))
    ec_paras = [p.text.strip() for p in ec_doc.paragraphs]
    ec_out = {"source": "edgecases-v1", "cases": []}
    for ec_id in sorted(EC_PARA_INDICES):
        indices, fw = EC_PARA_INDICES[ec_id]
        content = " ".join(ec_paras[i] for i in indices if ec_paras[i])
        ec_out["cases"].append({
            "id": f"edgecases-v1/{ec_id}",
            "framework": fw,
            "paragraphs": [content],
            "expect": {"badge": EC_EXPECTED[ec_id]},
        })
    corpus_files["edgecases-v1.json"] = ec_out

    # Paste demo (committed) — golden-only parity sample, no expectations
    demo_path = os.path.join(TEST_DOCS, "handoff", "Legal_Writing_Structure_Coach_Paste_Demo.txt")
    with open(demo_path, encoding="utf-8") as fh:
        demo_paras, _ = parse_plaintext(fh.read())
    corpus_files["paste-demo.json"] = {
        "source": "paste-demo",
        "cases": [{
            "id": "paste-demo/full",
            "framework": "AUTO",
            "paragraphs": demo_paras,
            "expect": {},
        }],
    }

    # Professor materials → corpus-local (gitignored)
    local_files = {}
    ma_path = os.path.join(ROOT, "model_answer_test_bank_v2.docx")
    if os.path.exists(ma_path):
        ma_doc = Document(ma_path)
        ma_paras = [p.text.strip() for p in ma_doc.paragraphs]
        ma_out = {"source": "model-answers-v2", "cases": []}
        for sec_id, indices, fw, badge in MODEL_ANSWER_SECTIONS:
            content = " ".join(ma_paras[i] for i in indices if i < len(ma_paras) and ma_paras[i])
            expect = {"badge": badge} if badge else {"badge_not_complete": True}
            if sec_id in MODEL_ANSWER_PENDING:
                expect["pending"] = MODEL_ANSWER_PENDING[sec_id]
            ma_out["cases"].append({
                "id": f"model-answers-v2/{sec_id}",
                "framework": fw,
                "paragraphs": [content],
                "expect": expect,
            })
        local_files["model-answers-v2.json"] = ma_out

    cp_path = os.path.join(ROOT, "CivPro_II_Exam1_Piedmont_Healthcare.docx")
    if os.path.exists(cp_path):
        cp_doc = Document(cp_path)
        cp_paras = [p.text.strip() for p in cp_doc.paragraphs]
        cp_out = {"source": "civpro-exam", "cases": []}
        for sec_id, indices, fw in CIVPRO_SECTIONS:
            para_texts = [cp_paras[i] for i in indices if i < len(cp_paras) and cp_paras[i]]
            cp_out["cases"].append({
                "id": f"civpro-exam/{sec_id}",
                "framework": fw,
                "paragraphs": para_texts,
                "expect": {
                    "para_expected_labels": [civpro_expected_label(p) for p in para_texts],
                },
            })
        local_files["civpro-exam.json"] = cp_out

    # Write corpus files
    for fname, data in corpus_files.items():
        path = os.path.join(corpus_dir, fname)
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, indent=2, ensure_ascii=False)
        print(f"wrote {path} ({len(data['cases'])} cases)")
    for fname, data in local_files.items():
        path = os.path.join(local_dir, fname)
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, indent=2, ensure_ascii=False)
        print(f"wrote {path} ({len(data['cases'])} cases)")

    # Golden runs. Committed golden covers ONLY the committed corpus; the
    # professor-material golden stays in the gitignored corpus-local dir.
    def run_golden(file_map, out_path):
        golden = {}
        n_sent = 0
        for data in file_map.values():
            for case in data["cases"]:
                result = classify_text(case["paragraphs"], case["framework"])
                golden[case["id"]] = result
                n_sent += result["summary"]["total_sentences"]
        with open(out_path, "w", encoding="utf-8") as fh:
            json.dump(golden, fh, indent=2, ensure_ascii=False)
        print(f"wrote {out_path} ({len(golden)} cases, {n_sent} sentences)")

    run_golden(corpus_files, os.path.join(golden_dir, "python-golden.json"))
    if local_files:
        run_golden(local_files, os.path.join(local_dir, "python-golden-local.json"))


if __name__ == "__main__":
    main()
