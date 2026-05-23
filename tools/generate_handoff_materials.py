from pathlib import Path
import html
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "handoff"


DEMO_SECTIONS = [
    (
        "1. Complete IRAC with ordinary signposts",
        "IRAC",
        "The issue is whether Mercer committed fraud. "
        "To establish fraud, a plaintiff must prove that the defendant made a false representation, knew it was false, intended reliance, and caused damages. "
        "Here, Mercer told Okafor that the investment fund returned twelve percent annually, but Mercer knew the fund had posted losses in four of the last ten years. "
        "Therefore, Okafor is likely to prevail on the fraud claim.",
        "Expected: Complete IRAC. Shows ordinary Issue, Rule, Application, Conclusion form.",
    ),
    (
        "2. Same IRAC with signposts removed",
        "IRAC",
        "The issue is whether Mercer committed fraud. "
        "To establish fraud, a plaintiff must prove that the defendant made a false representation, knew it was false, intended reliance, and caused damages. "
        "Mercer told Okafor that the investment fund returned twelve percent annually, but Mercer knew the fund had posted losses in four of the last ten years. "
        "Okafor is likely to prevail on the fraud claim.",
        "Expected: Complete IRAC. Demonstrates resilience when 'Here,' and 'Therefore,' are omitted.",
    ),
    (
        "3. Missing rule",
        "IRAC",
        "The issue is whether Park is liable for battery. "
        "Here, Park grabbed Lee's wrist during the argument and would not let go until a security guard intervened. "
        "Therefore, Lee is likely to establish battery.",
        "Expected: Missing Rule. The paragraph applies facts and concludes without first stating the governing rule.",
    ),
    (
        "4. Missing application",
        "IRAC",
        "The issue is whether the agreement is enforceable. "
        "A valid contract requires offer, acceptance, and consideration. "
        "Courts assess mutual assent under an objective theory of contracts.",
        "Expected: Missing Application. The paragraph states rules but never connects them to the facts.",
    ),
    (
        "5. Application before rule",
        "IRAC",
        "Harbor filed its personal jurisdiction objection in an answer after first moving to dismiss on venue grounds. "
        "A personal jurisdiction defense is waived if omitted from the first Rule 12 motion. "
        "Therefore, Harbor likely waived the defense.",
        "Expected: Application Before Rule. Useful for showing students why order matters.",
    ),
    (
        "6. Rule/Application blend",
        "IRAC",
        "The issue is whether Rivera acted negligently. "
        "A driver must use reasonable care, and here Rivera looked down at a text message for six seconds while entering the intersection. "
        "Therefore, Rivera likely breached the duty of care.",
        "Expected: Rule/Application Blend. One sentence mixes the abstract standard and the specific facts.",
    ),
    (
        "7. Complete CREAC",
        "CREAC",
        "The restrictive covenant is likely enforceable against Garrison. "
        "An equitable servitude binds later owners when the covenant touches and concerns land, the original parties intended it to run, and the later owner had notice. "
        "In Neponsit, the court enforced a covenant requiring property owners to pay fees because the promise affected the land and bound successors with notice. "
        "In this case, the covenant limits Parcel 14 to residential use, was recorded in the chain of title, and was reviewed by Garrison before purchase. "
        "Accordingly, the covenant is likely enforceable against Garrison.",
        "Expected: Complete CREAC. Shows opening conclusion, rule, explanation, application, closing conclusion.",
    ),
    (
        "8. Citation-heavy application without Here",
        "IRAC",
        "The issue is whether transfer was proper. "
        "Under 28 U.S.C. § 1404(a), a court may transfer for convenience when venue is proper. "
        "The case was transferred under § 1406(a), but the original venue was proper under § 1391. "
        "The transfer was improper.",
        "Expected: Complete IRAC. Demonstrates the new scored-evidence resilience around statutory citations.",
    ),
    (
        "9. Ambiguous structure for confidence coaching",
        "AUTO",
        "The court should consider fairness and efficiency.",
        "Expected: Introductory/Transition with low or medium confidence. Shows that ambiguous sentences are now treated as coaching prompts rather than certain answers.",
    ),
    (
        "10. Missing rule with revision priorities",
        "AUTO",
        "Here, Dana shoved Lee during the argument. "
        "Therefore, Dana is likely liable.",
        "Expected: Missing Rule. Shows paragraph-level structure score, revision priorities, and sentence-level next-step hints.",
    ),
]


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(30, 64, 175)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_handoff() -> Path:
    path = OUT / "IRAC_CREAC_Highlighter_Handoff.docx"
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_heading("IRAC/CREAC Structural Highlighter: Handoff", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Offline form-feedback tool for 1L legal writing and practice exam answers")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    add_heading(doc, "Purpose", 1)
    doc.add_paragraph(
        "This project is a form and structure checker, not a legal merits checker. "
        "It helps students see whether a paragraph contains the core moves of IRAC, RAC, CRAC, or CREAC: issue/position, rule, explanation where applicable, application, and conclusion. "
        "The May 2026 refinement turns the app from a pure highlighter into a 1L structure-coaching surface with confidence, revision priorities, and plain-English next steps."
    )

    add_heading(doc, "Why It Is Useful", 1)
    add_bullets(doc, [
        "Gives immediate visual feedback while students are learning legal writing form.",
        "Works offline and can be shared without sending student drafts to a hosted AI service.",
        "Focuses on structural habits that 1Ls often struggle to internalize before midterms and memo deadlines.",
        "Can support office hours, ILS workshops, TA sessions, and practice-exam self-review.",
    ])

    add_heading(doc, "Current Build", 1)
    add_bullets(doc, [
        "Inputs: paste text, .docx upload, .pdf upload, and .rtf upload.",
        "Frameworks: automatic per-paragraph detection across IRAC, RAC, CRAC, and CREAC.",
        "Analysis: sentence-level labels, paragraph badges, structural warnings, suggestions, cross-paragraph notes, and paragraph-level structure scores.",
        "Training layer: paragraph coaching panels summarize the first revision priority and list concrete fixes such as missing rule, missing application, blended rule/application, premature application, or low-confidence highlights.",
        "Evidence scoring: each sentence carries label scores, competing labels, selected evidence, high/medium/low confidence, uncertainty language, and a short revision hint.",
        "UI: structured hover tooltips explain why a sentence was classified, why confidence is not high, other plausible labels, and what to check next.",
        "Focus workflow: a filter can show only paragraphs with warnings, blends, or low-confidence sentences.",
        "Exports: annotated .docx with highlights, paragraph badges, average confidence, structure score, training summary, and top revision priorities.",
    ])

    add_heading(doc, "1L Training Behavior", 1)
    doc.add_paragraph(
        "The classifier remains deterministic and offline, but the student-facing experience now frames the output as coaching rather than a final grade. "
        "The highlighted label is still the app's selected label, while confidence and competing-label information make uncertainty visible."
    )
    add_bullets(doc, [
        "High confidence means multiple or strong structural signals support the label.",
        "Medium confidence means the label is plausible but depends on context or a narrower signal.",
        "Low confidence means the sentence should be treated as a prompt for human review, not a final answer.",
        "Revision priorities are deliberately structure-focused; they do not judge legal correctness, doctrinal accuracy, grammar, or style except where form affects IRAC/CREAC structure.",
        "Structure scores are rough completeness signals from 0 to 100, useful for triage and discussion rather than grading.",
    ])

    add_heading(doc, "Recommended Demo Script", 1)
    add_bullets(doc, [
        "Paste the first two demo examples to show that removing 'Here,' and 'Therefore,' no longer breaks the analysis.",
        "Upload the demo DOCX to show automatic IRAC/RAC/CRAC/CREAC detection, warning categories, structure scores, and revision priorities in one document.",
        "Upload the demo PDF to show that text-based PDFs follow the same analysis path.",
        "Hover over highlighted sentences to show the structured tooltip: why this label, other plausible labels, confidence/uncertainty, and what to check next.",
        "Use example 9 to show that ambiguous writing is treated cautiously instead of being overclaimed as certain.",
        "Use example 10 to show the coaching panel and the warnings/low-confidence-only focus filter.",
        "Export an annotated DOCX to show what a student could bring to office hours, including structure score and top revision priorities.",
    ])

    add_heading(doc, "Known Limits", 1)
    add_bullets(doc, [
        "The tool does not know whether the legal rule is correct.",
        "The tool does not evaluate the strength of the argument or factual inferences.",
        "Scanned PDFs still require OCR before this app can analyze them.",
        "Very unusual professor-specific formats may require custom examples and calibration.",
        "Confidence is heuristic calibration, not statistical model probability.",
        "Structure scores are not grades and should not be presented as assessment outcomes.",
    ])

    add_heading(doc, "Potential Added Features", 1)
    add_bullets(doc, [
        "Professor mode: let an instructor define preferred labels, colors, and warning language.",
        "Rubric mode: produce a printable structural checklist for each paragraph.",
        "Revision mode: compare before/after drafts and show whether structure improved.",
        "Practice exam mode: identify answer subparts and flag paragraphs that skip application.",
        "Calibration set: store professor-approved sample answers and tune phrase weights against them.",
        "Student reflection prompts: after warnings, ask the student to rewrite the missing rule/application/conclusion.",
        "Privacy-first packaged app: one-click executable or local web bundle for nontechnical users.",
    ])

    add_heading(doc, "Validation Snapshot", 1)
    add_bullets(doc, [
        "Baseline classifier demo tests: python test_classifier.py.",
        "Structural spectrum tests: python test_structural_spectrum.py.",
        "Perturbation tests: python test_perturbations.py.",
        "Training-layer regression tests: python test_training_tool.py.",
        "JS syntax check: bundled Node runtime with node --check app/static/app.js.",
        "Live smoke path: /analyze returns training fields and /export produces a valid annotated DOCX.",
    ])

    add_heading(doc, "Sharing Options", 1)
    add_bullets(doc, [
        "Simplest feedback path: zip the project folder and include this handoff plus the demo files. The recipient installs Python, runs pip install -r requirements.txt, then runs python run.py.",
        "More polished path: build a one-click desktop executable with PyInstaller or package it as a local app. This is better for professor feedback because it avoids command-line setup.",
        "Web feedback path: deploy a private demo instance, but only if you are comfortable with draft text being uploaded to that server. For student privacy, the local/offline version is preferable.",
        "GitHub path: put the project in a private GitHub repository and invite the professor as a collaborator or viewer. This is best if you want issue-based feedback and version history.",
    ])

    doc.save(path)
    return path


def build_demo_docx() -> Path:
    path = OUT / "IRAC_CREAC_Demo_Test_Document.docx"
    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_heading("IRAC/CREAC Demo Test Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Use this document to test DOCX upload. The same examples are also provided as paste text and as a text-based PDF. "
        "The updated examples exercise the May 2026 training-tool behavior: structure scores, revision priorities, confidence, competing labels, and richer hover/export explanations."
    )
    for heading, framework, text, expected in DEMO_SECTIONS:
        add_heading(doc, heading, 2)
        p = doc.add_paragraph()
        p.add_run(f"Framework: {framework}. ").bold = True
        p.add_run(expected)
        doc.add_paragraph(text)
    doc.save(path)
    return path


def build_paste_text() -> Path:
    path = OUT / "IRAC_CREAC_Paste_Demo.txt"
    parts = []
    for heading, framework, text, expected in DEMO_SECTIONS:
        parts.append(f"{heading}\nFramework: {framework}. {expected}\n{text}")
    path.write_text("\n\n".join(parts), encoding="utf-8")
    return path


def build_local_readme() -> Path:
    path = OUT / "LOCAL_BUILD_README.txt"
    path.write_text(
        """IRAC/CREAC Structural Highlighter - Local Build

How to run
1. Open the IRAC_CREAC_Highlighter folder.
2. Double-click IRAC_CREAC_Highlighter.exe.
3. A browser window should open at a local address such as http://localhost:5050.
4. Paste text or upload a DOCX/PDF/RTF document.
5. Hover over highlighted sentences for the structured explanation.
6. Use Export Annotated DOCX to create a highlighted review copy with structure score and revision priorities.

What this tool checks
- Legal writing form and structure only.
- Automatic paragraph-level detection for IRAC, RAC, CRAC, and CREAC.
- Sentence labels, paragraph-level completion/warning badges, structure scores, confidence, competing labels, and revision-priority coaching.
- Training-oriented warnings such as missing rule, missing application, missing conclusion, blended rule/application, premature application, and low-confidence highlights.

What this tool does not check
- Whether the legal rule is correct.
- Whether the answer is doctrinally persuasive.
- Whether the student would earn a particular grade.
- Whether a structure score should be treated as a formal assessment.

Demo files
- IRAC_CREAC_Highlighter_Handoff.docx explains the project, current training-tool behavior, validation, and suggested uses.
- IRAC_CREAC_Demo_Test_Document.docx tests DOCX upload with clean, missing, blended, citation-heavy, ambiguous, and coaching-priority examples.
- IRAC_CREAC_Demo_Test_Document.pdf tests text-based PDF upload.
- IRAC_CREAC_Paste_Demo.txt contains examples for paste testing.
- IRAC_CREAC_1L_Structural_Spectrum_Test.docx stresses clean, missing, muddled, and citation-heavy examples.
- CHANGELOG.md summarizes the May 2026 training-tool refinement.

Privacy note
This build runs locally on the computer. It is preferable for student drafts because the text is not sent to a hosted web service.
""",
        encoding="utf-8",
    )
    return path


def build_changelog() -> Path:
    path = ROOT / "CHANGELOG.md"
    path.write_text(
        """# Changelog

## 2026-05-22 - 1L Training Tool Refinement

- Added a training-oriented coaching layer while keeping the app fully offline and deterministic.
- Added sentence-level diagnostics: confidence, competing labels, uncertainty reasons, evidence, trigger phrase, and revision hints.
- Added paragraph-level coaching: structure score, training summary, and ordered revision priorities.
- Added UI support for coaching panels, structured hover tooltips, confidence legend, and a warnings/low-confidence-only focus filter.
- Expanded DOCX export to include average confidence, structure score, training summary, and top revision priorities.
- Cleaned up duplicated classifier helper definitions.
- Added `test_training_tool.py` for the training/coaching response fields and cross-paragraph priority behavior.
- Refreshed demo and handoff materials to reflect the new coaching workflow.

Validation:
- `python test_classifier.py`
- `python test_structural_spectrum.py`
- `python test_perturbations.py`
- `python test_training_tool.py`
- Bundled Node syntax check for `app/static/app.js`
- Live `/analyze` and `/export` smoke test

## 2026-04-27 - Local Demo Build and Handoff Materials

- Generated local build handoff materials for professor/demo review.
- Added demo DOCX, text-based PDF, paste-demo text, local build README, and 1L structural spectrum test document.
- Packaged the app for local/offline review.
""",
        encoding="utf-8",
    )
    return path


def build_demo_pdf() -> Path:
    path = OUT / "IRAC_CREAC_Demo_Test_Document.pdf"
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="SmallNote", parent=styles["BodyText"], fontSize=8.5, leading=11, textColor="#374151"))
        story = [
            Paragraph("IRAC/CREAC Demo Test Document", styles["Title"]),
            Paragraph("Use this text-based PDF to test PDF upload.", styles["SmallNote"]),
            Spacer(1, 0.2 * inch),
        ]
        for heading, framework, text, expected in DEMO_SECTIONS:
            story.append(Paragraph(heading, styles["Heading2"]))
            story.append(Paragraph(f"<b>Framework: {framework}.</b> {expected}", styles["SmallNote"]))
            story.append(Paragraph(text, styles["BodyText"]))
            story.append(Spacer(1, 0.12 * inch))
        pdf = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
        pdf.build(story)
    except ModuleNotFoundError:
        chrome = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")
        if not chrome.exists():
            print("Skipping demo PDF: reportlab is not installed and Chrome fallback was not found.")
            return path
        html_path = OUT / "IRAC_CREAC_Demo_Test_Document.html"
        sections = []
        for heading, framework, text, expected in DEMO_SECTIONS:
            sections.append(
                f"<h2>{html.escape(heading)}</h2>"
                f"<p class='note'><strong>Framework: {html.escape(framework)}.</strong> {html.escape(expected)}</p>"
                f"<p>{html.escape(text)}</p>"
            )
        html_path.write_text(
            "<!doctype html><html><head><meta charset='utf-8'>"
            "<style>body{font-family:Arial,sans-serif;margin:42px;line-height:1.45;color:#111827}"
            "h1{font-size:24px;margin:0 0 6px}h2{font-size:15px;margin:18px 0 4px;color:#1e40af}"
            "p{font-size:11px;margin:4px 0}.note{font-size:10px;color:#374151}</style></head><body>"
            "<h1>IRAC/CREAC Demo Test Document</h1>"
            "<p class='note'>Use this text-based PDF to test PDF upload.</p>"
            + "".join(sections)
            + "</body></html>",
            encoding="utf-8",
        )
        subprocess.run([
            str(chrome),
            "--headless=new",
            "--disable-gpu",
            f"--print-to-pdf={path}",
            str(html_path),
        ], check=True)
    return path


def main() -> None:
    OUT.mkdir(exist_ok=True)
    paths = [
        build_handoff(),
        build_demo_docx(),
        build_paste_text(),
        build_demo_pdf(),
        build_local_readme(),
        build_changelog(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
