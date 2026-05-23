from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from test_structural_spectrum import CASES

OUT = ROOT / "handoff" / "IRAC_CREAC_1L_Structural_Spectrum_Test.docx"


def main() -> None:
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("IRAC/CREAC 1L Structural Spectrum Test", level=0)
    doc.add_paragraph(
        "This document is designed to exercise automatic format detection across clean IRAC, RAC, CRAC, CREAC, missing components, blended sentences, and citation-heavy application."
    )

    for idx, case in enumerate(CASES, start=1):
        doc.add_heading(f"{idx}. {case['name']}", level=1)
        meta = doc.add_paragraph()
        meta.add_run("Expected badge: ").bold = True
        meta.add_run(case["badge"])
        if case.get("effective"):
            meta.add_run(" | Expected detected format: ").bold = True
            meta.add_run(case["effective"])
        doc.add_paragraph(case["text"])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
